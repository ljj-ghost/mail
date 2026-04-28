from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DESKTOP = Path.home() / "Desktop"
OUTPUT = DESKTOP / "商城微服务项目技术文档-小白版.docx"
PROJECT_ROOT = Path(r"D:\ideaproject\mail")

ACCENT = "2F5D62"
ACCENT_2 = "526D82"
LIGHT = "EEF5F6"
LIGHT_2 = "F7F9FA"
WARN = "FFF4DB"
TEXT = "1F2933"
MUTED = "667085"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill: str = ACCENT, header_text: str = "FFFFFF") -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_text)


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]) -> None:
    headers = list(headers)
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    style_table(table)
    doc.add_paragraph()


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line_index, line in enumerate(code.strip("\n").splitlines()):
        if line_index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("344054")
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(ACCENT if level <= 2 else ACCENT_2)
        run.bold = True


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)

    for name, size in [("Title", 26), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.font.bold = True

    return doc


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("商城微服务项目技术文档")
    r.bold = True
    r.font.size = Pt(28)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("小白友好版：从技术栈到业务逻辑逐步讲解")
    r.font.size = Pt(14)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    add_callout(
        doc,
        "阅读定位",
        "这不是只给资深工程师看的架构说明，而是一份带你认识项目的入门文档。它会先解释技术名词，再解释每个模块为什么存在，最后把登录、加购、下单、支付、扣库存这些流程串起来。",
    )

    info = doc.add_table(rows=5, cols=2)
    rows = [
        ("项目路径", str(PROJECT_ROOT)),
        ("文档生成位置", str(OUTPUT)),
        ("项目类型", "在线商城微服务项目，包含 Java 后端、API 网关、React 前端和本地部署脚本"),
        ("主要读者", "刚接触 Java / Spring Boot / React / 微服务的新手"),
        ("当前代码状态", "开发基线版本：核心认证、商品、购物车、订单、支付、库存流程可读可跑，Nacos 与 RocketMQ 主要作为后续扩展预留"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_text(info.rows[i].cells[0], k, bold=True, color="FFFFFF")
        set_cell_text(info.rows[i].cells[1], v)
        set_cell_shading(info.rows[i].cells[0], ACCENT)
        set_cell_margins(info.rows[i].cells[0])
        set_cell_margins(info.rows[i].cells[1])
    doc.add_page_break()


def build_doc() -> Document:
    doc = setup_document()
    add_cover(doc)

    add_heading(doc, "1. 先把项目看成一张地图", 1)
    add_para(
        doc,
        "这个项目可以理解成一个小型在线商城。用户从前端页面进入系统，先经过 API 网关；网关确认用户身份后，把请求转发给对应的后端服务。每个服务只负责一块业务：认证管登录，商品管商品，购物车管加购，订单管下单，支付管支付单，库存管锁库存和扣库存。",
    )
    add_callout(
        doc,
        "一句话理解微服务",
        "微服务不是把代码随便拆成很多文件，而是把不同业务边界拆成独立服务。这样商品服务出问题时，不应该直接影响认证服务；订单服务需要商品信息时，通过接口去问商品服务，而不是直接读商品服务的表。",
    )
    add_table(
        doc,
        ["模块", "端口", "负责什么", "新手理解"],
        [
            ("mall-gateway", "18080", "统一入口、路由转发、鉴权、透传用户信息", "像商场入口保安，先看你有没有通行证，再告诉你去几楼"),
            ("auth-service", "18081", "登录、注册、刷新令牌、登出、解析令牌、管理员用户管理", "负责证明“你是谁”"),
            ("user-service", "18082", "用户资料、收货地址", "负责用户档案"),
            ("product-service", "18083", "商品分类、商品列表、详情、后台商品维护、图片上传", "负责货架和商品说明书"),
            ("inventory-service", "18084", "库存查询、预占、释放、正式扣减、低库存预警", "负责仓库数量"),
            ("cart-service", "18085", "购物车列表、加入购物车、删除、下单后清理", "负责用户临时挑选的商品"),
            ("order-service", "18086", "提交订单、订单列表、取消、发货、完成、支付回写", "负责交易状态机"),
            ("payment-service", "18087", "创建支付单、查询支付单、模拟支付成功、关闭支付单", "负责收银台"),
            ("mall-web", "Vite 默认端口", "React 前端页面", "用户真正看到和点击的商城界面"),
        ],
    )

    add_heading(doc, "2. 技术栈逐个解释", 1)
    add_para(doc, "下面这些技术名词在项目里反复出现。先把它们讲清楚，后面读代码就不会像看一堆缩写。")
    add_table(
        doc,
        ["技术", "项目中怎么用", "新手理解"],
        [
            ("Java 21", "后端服务的开发语言，根 pom.xml 里设置 java.version=21", "Java 是后端主语言，21 是版本号"),
            ("Maven 多模块", "根工程聚合 mall-common、mall-api、mall-gateway、mall-services 等模块", "像一个大文件夹管理多个子项目，并统一构建"),
            ("Spring Boot 3.2.4", "每个服务都是可独立启动的 Spring Boot 应用", "帮你快速启动 Web 服务、加载配置、管理 Bean"),
            ("Spring Cloud Gateway", "mall-gateway 用它做路由和鉴权过滤", "统一入口，把 /api/v1/products 转给商品服务"),
            ("OpenFeign", "订单调库存、商品、购物车；支付调订单；购物车调商品", "像写 Java 接口一样调用另一个服务的 HTTP 接口"),
            ("JdbcTemplate", "Repository 直接写 SQL 操作 MySQL", "比 ORM 更直观：SQL 写什么，数据库就执行什么"),
            ("MySQL 8", "保存账号、商品、购物车、库存、订单、支付数据", "系统的长期记忆"),
            ("Redis 7", "缓存商品、用户、购物车、库存、订单详情，也做下单幂等锁", "系统的短期记忆，速度快但不能当唯一真相"),
            ("JWT / JJWT", "auth-service 签发 access token 和 refresh token", "登录成功后拿到的电子通行证"),
            ("BCrypt", "登录密码不明文保存，而是保存加密哈希", "即使数据库泄露，也不能直接看到用户密码"),
            ("Jakarta Validation", "用 @Valid、@Min 等检查请求参数", "入口处先检查数据是不是合格"),
            ("Docker Compose", "deploy 目录一键启动 MySQL、Redis、Nacos、RocketMQ", "给开发环境拉起基础设施"),
            ("React 19 + TypeScript", "mall-web 前端页面和交互逻辑", "React 负责画页面，TypeScript 帮你提前发现类型错误"),
            ("Vite", "前端开发和打包工具", "让前端本地启动更快"),
            ("React Router", "控制 /products、/cart、/checkout 等页面跳转", "前端页面的导航系统"),
        ],
    )
    add_callout(
        doc,
        "注意当前项目的真实情况",
        "mall-dependencies 统一管理了 MyBatis-Plus、MapStruct、Knife4j、RocketMQ、Nacos 等版本或配置方向，但当前核心代码主要使用 JdbcTemplate、OpenFeign、Redis 和同步 HTTP 调用。RocketMQ 与 Nacos 已在配置和部署层预留，业务主链路还没有完全改成 MQ 事件驱动。",
        WARN,
    )

    add_heading(doc, "3. 目录结构怎么读", 1)
    add_code(
        doc,
        r"""
D:\ideaproject\mail
├─ pom.xml                         根 Maven 工程，只聚合模块
├─ mall-dependencies               统一依赖版本和插件版本
├─ mall-common                     公共能力：统一响应、异常、安全上下文
├─ mall-api                        服务之间共享的 DTO / Request / Response
├─ mall-gateway                    API 网关，外部请求统一从这里进来
├─ mall-services                   业务微服务集合
│  ├─ auth-service                 认证服务
│  ├─ user-service                 用户服务
│  ├─ product-service              商品服务
│  ├─ inventory-service            库存服务
│  ├─ cart-service                 购物车服务
│  ├─ order-service                订单服务
│  └─ payment-service              支付服务
├─ mall-web                        React + Vite 前端
├─ deploy                          本地中间件和启动脚本
└─ docs                            设计文档、接口清单、SQL 设计资料
""",
    )
    add_bullets(
        doc,
        [
            "mall-common 不是业务服务，它提供所有服务都要用的基础能力，例如统一返回体 CommonResponse、BusinessException、GlobalExceptionHandler、UserContext。",
            "mall-api 也不放业务逻辑，它只放跨服务通信时要共同理解的数据结构。例如订单服务调用库存服务时，双方都要认识 InventoryReserveRequest。",
            "mall-services 才是真正写业务的地方。每个服务都有 controller、service、repository、resources/db/schema.sql 这一套结构。",
            "mall-web 是独立的前端项目，使用 npm 管理依赖，不参与 Maven 构建。",
        ],
    )

    add_heading(doc, "4. Maven 多模块逻辑", 1)
    add_para(
        doc,
        "Maven 的作用是管理依赖、编译、打包和模块关系。这个项目根目录的 pom.xml 不写业务代码，只声明有哪些子模块。mall-dependencies 进一步统一锁定 Spring Boot、Spring Cloud、JJWT、Lombok、Knife4j 等版本，避免每个子模块自己乱配版本。",
    )
    add_table(
        doc,
        ["Maven 模块", "打包类型", "职责"],
        [
            ("mall-parent", "pom", "根聚合工程，声明 Java 21 和所有一级模块"),
            ("mall-dependencies", "pom", "统一版本管理，服务模块继承它"),
            ("mall-common", "pom", "聚合 common-core、common-web、common-security"),
            ("mall-api", "pom", "聚合各业务域共享 DTO 模块"),
            ("mall-gateway", "jar", "可启动网关应用"),
            ("mall-services/*-service", "jar", "每个业务服务都是可启动 Spring Boot 应用"),
        ],
    )
    add_code(
        doc,
        r"""
# 全量打包，跳过测试
& 'D:\apache-maven-3.6.3\bin\mvn.cmd' -q -DskipTests package

# 只打包订单服务，同时自动带上它依赖的模块
& 'D:\apache-maven-3.6.3\bin\mvn.cmd' -pl mall-services/order-service -am -DskipTests package
""",
    )

    add_heading(doc, "5. 后端通用请求处理模型", 1)
    add_para(doc, "大多数后端服务都遵循同一种代码路径：Controller 接请求，Service 做业务判断和流程编排，Repository 写 SQL，返回 CommonResponse。")
    add_numbered(
        doc,
        [
            "前端或其他服务发起 HTTP 请求。",
            "Controller 用 @GetMapping、@PostMapping 接收请求，用 @Valid 校验参数。",
            "Controller 调用 Service，不在 Controller 里写复杂业务。",
            "Service 判断用户身份、业务状态、库存是否足够、是否允许操作。",
            "Service 调用 Repository 写入或查询 MySQL，必要时读写 Redis 缓存。",
            "如果出现业务错误，抛 BusinessException。",
            "GlobalExceptionHandler 把异常转换成统一 JSON：code、message、data、traceId。",
        ],
    )
    add_table(
        doc,
        ["类/机制", "路径", "作用"],
        [
            ("CommonResponse", "mall-common-core", "统一返回格式：成功 code=0，失败带错误码和 traceId"),
            ("BusinessException", "mall-common-core", "表示可预期的业务错误，例如库存不足、无权限"),
            ("GlobalExceptionHandler", "mall-common-web", "拦截异常，转换成标准响应"),
            ("UserContext", "mall-common-security", "用 ThreadLocal 保存当前请求用户 ID 和角色"),
            ("MockUserHeaderFilter", "mall-common-security", "从 X-User-Id、X-User-Role 请求头读取用户信息"),
        ],
    )

    add_heading(doc, "6. 网关：请求从哪里进来", 1)
    add_para(
        doc,
        "mall-gateway 是系统的统一入口。外部前端通常访问 18080 端口，网关根据路径决定把请求转发给哪个服务。例如 /api/v1/products/** 转发到 product-service，/api/v1/orders/** 转发到 order-service。",
    )
    add_table(
        doc,
        ["路径", "转发目标", "说明"],
        [
            ("/api/v1/auth/**", "auth-service:18081", "登录、注册、刷新令牌"),
            ("/api/v1/users/**", "user-service:18082", "用户资料和地址"),
            ("/api/v1/products/**", "product-service:18083", "商品浏览、图片资源"),
            ("/api/v1/cart/**", "cart-service:18085", "购物车"),
            ("/api/v1/orders/**", "order-service:18086", "订单"),
            ("/api/v1/payments/**", "payment-service:18087", "支付"),
        ],
    )
    add_para(doc, "网关的鉴权逻辑在 GatewayAuthenticationFilter 中：")
    add_numbered(
        doc,
        [
            "OPTIONS、/actuator、登录、注册、商品浏览等公开路径直接放行。",
            "internal 开头的内部接口禁止从网关暴露，防止外部直接调用服务内部能力。",
            "受保护的 /api 路径必须携带 Authorization: Bearer <accessToken>。",
            "网关调用 auth-service 的 /internal/v1/auth/tokens/parse 校验 token。",
            "校验成功后，网关把 X-User-Id、X-User-Role、X-Session-No 等请求头传给下游服务。",
        ],
    )

    add_heading(doc, "7. 认证服务：登录为什么能记住你", 1)
    add_para(
        doc,
        "auth-service 负责账号、密码、会话和 JWT。登录成功后，服务返回 accessToken 和 refreshToken。前端把它们保存到 localStorage，之后访问需要登录的接口时带上 accessToken。",
    )
    add_table(
        doc,
        ["表", "作用", "关键字段"],
        [
            ("auth_account", "保存登录账号和密码哈希", "login_name、password_hash、status、role_code、fail_count"),
            ("auth_token_session", "保存一次登录会话", "session_no、access_jti、refresh_jti、expire_time、status"),
            ("auth_login_log", "记录登录成功或失败", "login_name、client_ip、login_result、fail_reason"),
            ("mall_user.ums_user", "用户基础资料", "user_id、nickname、mobile、email、role_code"),
        ],
    )
    add_heading(doc, "7.1 登录流程", 2)
    add_numbered(
        doc,
        [
            "AuthController 接收 /api/v1/auth/login/password。",
            "AuthService 根据 loginName 查 auth_account。",
            "如果账号不存在、禁用或失败次数过多，记录失败日志并返回业务错误。",
            "用 BCrypt 校验用户输入密码和数据库里的 password_hash。",
            "校验成功后，JwtTokenService 生成 accessToken 和 refreshToken。",
            "把 session_no、access_jti、refresh_jti 写入 auth_token_session。",
            "把会话快照缓存到 Redis：mall:auth:session:{sessionNo}。",
            "返回 TokenInfoDTO 给前端。",
        ],
    )
    add_heading(doc, "7.2 refreshToken 为什么存在", 2)
    add_para(
        doc,
        "accessToken 一般较短，用来访问接口；refreshToken 更长，用来换新 token。当前项目刷新时会校验 refresh token 的 tokenType 必须是 REFRESH，然后轮换 access_jti 和 refresh_jti。轮换的意思是旧 token 之后不能继续使用，降低泄露风险。",
    )
    add_heading(doc, "7.3 管理员权限", 2)
    add_para(doc, "用户角色通过 role_code 区分 USER 和 ADMIN。网关解析 token 后把 X-User-Role 传给服务，服务里的 UserContext.requireAdmin() 会检查当前用户是不是管理员。")

    add_heading(doc, "8. 用户服务：用户资料和地址", 1)
    add_para(
        doc,
        "user-service 管理用户资料和收货地址。它有两类接口：用户自己操作的接口，以及其他服务内部查询的接口。比如订单服务下单时，需要知道用户默认地址，就可以调用 /internal/v1/users/{userId}/default-address。",
    )
    add_bullets(
        doc,
        [
            "查询当前用户：/api/v1/users/me，返回 UserBaseDTO。",
            "更新资料：/api/v1/users/profile，更新 nickname、mobile、email。",
            "地址列表：/api/v1/users/addresses，按默认地址优先排序。",
            "新增地址时，如果用户还没有默认地址，新地址会自动成为默认地址。",
            "删除默认地址时，会把剩余地址中的第一条设为默认地址，避免用户没有默认地址。",
        ],
    )
    add_callout(
        doc,
        "缓存逻辑",
        "用户资料缓存 15 分钟，地址列表缓存 10 分钟。更新资料、增删改地址后会删除对应 Redis 缓存，下次查询再从 MySQL 读最新数据。",
    )

    add_heading(doc, "9. 商品服务：商城货架", 1)
    add_para(
        doc,
        "product-service 提供前台商品浏览和后台商品维护。前台关注分类、商品列表、商品详情；后台关注商品增删改、上下架、图片上传和商品统计。",
    )
    add_table(
        doc,
        ["表", "含义", "例子"],
        [
            ("pms_category", "商品分类", "手机、耳机、笔记本"),
            ("pms_spu", "标准产品单元，表示一个产品系列", "Nova X16"),
            ("pms_sku", "具体可购买商品，带价格、图片、库存对应编号", "Nova X16 256G 曜石黑"),
        ],
    )
    add_para(
        doc,
        "新手最容易混淆 SPU 和 SKU：SPU 是一类商品，SKU 是具体规格。用户真正下单的是 SKU，因为库存、价格、订单明细都要精确到具体规格。",
    )
    add_bullets(
        doc,
        [
            "recommend(limit)：返回首页推荐 SKU，缓存键 mall:product:recommend:{limit}。",
            "catalog(categoryId, keyword, limit)：按分类或关键词搜索商品，缓存键包含分类、关键词和数量。",
            "skuDetail(skuId)：加载商品详情，缓存 15 分钟。",
            "internalSku / batchSkus：给购物车和订单服务使用，只返回 skuId、spuId、skuName、salePrice、status 等最小数据。",
            "后台创建、更新、删除商品后，会扫描并删除 mall:product:* 相关缓存。",
        ],
    )
    add_callout(
        doc,
        "图片上传安全点",
        "ProductAssetService 只允许 JPG、PNG、WEBP、GIF，并且会用随机文件名保存到 deploy/uploads/product-images。读取图片时会校验文件名，防止通过 ../ 访问目录外文件。",
    )

    add_heading(doc, "10. 购物车服务：临时购物清单", 1)
    add_para(
        doc,
        "cart-service 的逻辑很直观：用户把 SKU 加入购物车，服务先问 product-service 这个 SKU 是否存在且可售，然后写入 oms_cart_item。相同用户、相同 SKU 使用唯一键，重复加入时数量累加。",
    )
    add_numbered(
        doc,
        [
            "前端调用 /api/v1/cart/items。",
            "CartService 从 UserContext 取当前用户 ID。",
            "通过 ProductClient 调 product-service 的内部 SKU 接口。",
            "SKU 不存在或 status 不为 1，直接返回业务错误。",
            "CartRepository 执行 INSERT ... ON DUPLICATE KEY UPDATE。",
            "删除 Redis 购物车缓存 mall:cart:items:{userId}。",
        ],
    )
    add_para(doc, "下单成功后，order-service 会调用 cart-service 的内部接口 /internal/v1/cart/items/clear，清理已经提交订单的 SKU。")

    add_heading(doc, "11. 库存服务：防止卖超", 1)
    add_para(
        doc,
        "inventory-service 是交易系统里最关键的服务之一。它不是简单地保存一个库存数字，而是区分 available_qty 和 locked_qty。可售库存 = available_qty - locked_qty。",
    )
    add_table(
        doc,
        ["字段/状态", "含义", "举例"],
        [
            ("available_qty", "真实可用库存", "仓库里原本有 100 件"),
            ("locked_qty", "已被订单预占但还没真正扣减的库存", "有 2 件被待支付订单锁住"),
            ("saleable_qty", "可继续卖给其他人的库存，代码里动态计算", "100 - 2 = 98"),
            ("预占状态 1", "RESERVED，订单待支付", "库存已锁住"),
            ("预占状态 2", "RELEASED，订单取消或超时", "锁住的库存还回可售池"),
            ("预占状态 3", "DEDUCTED，支付成功", "正式扣减真实库存"),
        ],
    )
    add_heading(doc, "11.1 预占库存流程", 2)
    add_numbered(
        doc,
        [
            "订单服务提交订单前调用库存 check，先看每个 SKU 是否够卖。",
            "库存服务检查 saleableQty 是否大于等于购买数量。",
            "订单服务生成 orderNo，再调用 reserve。",
            "reserve 先看该 orderNo 是否已经有预占记录，保证重复调用安全。",
            "对每个 SKU 执行条件更新：只有 available_qty - locked_qty >= quantity 才能加锁。",
            "插入 inventory_reservation 和 inventory_reservation_item。",
        ],
    )
    add_heading(doc, "11.2 释放和扣减", 2)
    add_bullets(
        doc,
        [
            "释放库存：订单取消或超时，locked_qty 减少，库存重新可卖。",
            "正式扣减：支付成功后，available_qty 和 locked_qty 同时减少，因为预占阶段已经减少了可售库存。",
            "重复释放已释放的预占，直接返回成功；重复扣减已扣减的预占，也直接返回成功。这就是幂等思想。",
        ],
    )

    add_heading(doc, "12. 订单服务：交易状态机", 1)
    add_para(
        doc,
        "order-service 是交易主线的大脑。它不直接管理商品价格来源，也不直接改支付表，但它负责订单状态流转，协调商品、库存、购物车、支付等服务。",
    )
    add_table(
        doc,
        ["订单状态码", "含义", "允许的主要动作"],
        [
            ("10", "待支付", "取消、删除、创建支付单、支付成功回写"),
            ("20", "待发货", "后台填写物流并发货"),
            ("30", "已发货/待收货", "用户确认收货"),
            ("40", "已完成", "正向流程结束"),
            ("50", "已取消", "不可继续支付，用户侧可软删除"),
        ],
    )
    add_table(
        doc,
        ["支付状态码", "含义"],
        [
            ("0", "待支付"),
            ("2", "支付成功"),
            ("4", "支付关闭"),
        ],
    )
    add_heading(doc, "12.1 提交订单完整流程", 2)
    add_numbered(
        doc,
        [
            "OrderController 接收 /api/v1/orders/submit。",
            "OrderService 从 UserContext 获取当前用户。",
            "读取 idempotencyKey，并在 Redis 设置 IN_PROGRESS，防止重复点击下单按钮。",
            "查询 oms_order_submit_log，如果这个幂等键已经生成过订单，直接返回原订单号。",
            "调用 product-service 批量获取 SKU 快照，校验商品存在且可售。",
            "调用 inventory-service check 校验库存是否足够。",
            "生成 orderNo，再调用 inventory-service reserve 预占库存。",
            "计算 payAmount，把 oms_order、oms_order_item、oms_order_submit_log 放在本地事务里写入。",
            "写入成功后，把幂等键对应的订单号缓存到 Redis 一天。",
            "调用 cart-service 清理本次下单的购物车商品。",
            "如果订单本地事务失败，会尽量调用 inventory-service release 释放刚刚预占的库存。",
        ],
    )
    add_callout(
        doc,
        "幂等到底是什么",
        "幂等就是同一个请求重复执行多次，最终结果仍然像执行一次。用户可能连续点两次提交订单，网络也可能重试。如果没有幂等，系统可能生成两个订单；当前项目用 Redis 短锁 + oms_order_submit_log 唯一记录来兜底。",
    )
    add_heading(doc, "12.2 取消订单", 2)
    add_para(
        doc,
        "只有待支付订单能取消。取消时先调用库存服务释放库存，再在订单库中把 order_status 改成 50、pay_status 改成 4，然后尝试通知支付服务关闭对应支付单。",
    )
    add_heading(doc, "12.3 支付成功回写订单", 2)
    add_para(
        doc,
        "支付服务模拟支付成功后，会调用 order-service 的 /internal/v1/orders/{orderNo}/paid。订单服务先让库存服务执行 confirm-deduct，扣减预占库存；扣库存成功后，订单状态改为待发货，支付状态改为成功。",
    )
    add_heading(doc, "12.4 超时关单", 2)
    add_para(
        doc,
        "当前项目提供 /internal/v1/orders/close-expired，按创建时间扫描超过指定分钟数仍待支付的订单，然后逐个执行取消逻辑。文档里规划了 RocketMQ 延迟消息，但当前代码里主要是内部接口扫描方式。",
    )

    add_heading(doc, "13. 支付服务：开发版收银台", 1)
    add_para(
        doc,
        "payment-service 当前实现的是开发演示版支付链路：创建支付单、返回一个 mockSuccessUrl，然后通过接口模拟第三方支付成功回调。它适合本地联调订单状态和库存扣减，但还不是完整的支付宝/微信支付对接。",
    )
    add_numbered(
        doc,
        [
            "前端调用 /api/v1/payments/create，传 orderNo 和 payChannel。",
            "PaymentService 调 order-service 查询订单基础信息。",
            "确认订单属于当前用户，且未支付、未关闭。",
            "如果这个订单已有未关闭支付单，直接返回原支付单，避免重复创建。",
            "没有支付单时，生成 PAY- 开头的 paymentNo，写入 pay_payment。",
            "返回 paymentNo、orderNo、payStatus、payAmount 和 mockSuccessUrl。",
            "前端点击模拟成功时，调用 /api/v1/payments/mock/success/{paymentNo}。",
            "支付服务先调用订单服务 markPaid，订单扣库存并更新订单状态。",
            "订单成功后，支付服务把 pay_payment 改成支付成功，写 thirdTradeNo 和 payTime。",
        ],
    )
    add_callout(
        doc,
        "为什么支付服务不直接改订单表",
        "微服务边界要求每个服务只改自己的数据。支付服务只改 pay_payment；订单状态必须由 order-service 自己改。支付服务通过内部接口告诉订单服务“这笔支付成功了”。",
    )

    add_heading(doc, "14. 核心交易流程总览", 1)
    add_heading(doc, "14.1 登录到访问受保护接口", 2)
    add_code(
        doc,
        """
用户输入账号密码
  -> auth-service 校验 BCrypt 密码
  -> 生成 accessToken / refreshToken
  -> 前端保存到 localStorage
  -> 前端请求订单接口时带 Authorization: Bearer accessToken
  -> mall-gateway 调 auth-service 解析 token
  -> mall-gateway 写入 X-User-Id / X-User-Role
  -> order-service 通过 UserContext 读取当前用户
""",
    )
    add_heading(doc, "14.2 浏览商品到加入购物车", 2)
    add_code(
        doc,
        """
前端商品页
  -> GET /api/v1/products 或 /api/v1/products/sku/{skuId}
  -> product-service 读 Redis 缓存
  -> 缓存没有则查 MySQL pms_sku / pms_spu / pms_category
  -> 用户点击加入购物车
  -> cart-service 调 product-service 校验 SKU
  -> 写 oms_cart_item，重复 SKU 则数量累加
  -> 删除购物车缓存
""",
    )
    add_heading(doc, "14.3 下单到支付成功", 2)
    add_code(
        doc,
        """
提交订单
  -> order-service 做幂等检查
  -> product-service 返回 SKU 快照和价格
  -> inventory-service 检查库存
  -> inventory-service 预占库存 locked_qty += quantity
  -> order-service 写订单和订单项
  -> cart-service 清理已下单 SKU
  -> payment-service 创建支付单
  -> payment-service 模拟支付成功
  -> order-service 调 inventory-service 正式扣减
  -> order-service 改订单为待发货
  -> payment-service 改支付单为成功
""",
    )
    add_heading(doc, "14.4 取消订单", 2)
    add_code(
        doc,
        """
用户取消待支付订单
  -> order-service 校验订单属于当前用户
  -> inventory-service release 释放预占库存
  -> order-service 将订单状态改成已取消、支付状态改成已关闭
  -> payment-service closeByOrder 关闭未支付的支付单
  -> 删除订单详情缓存
""",
    )

    add_heading(doc, "15. Redis 在项目里承担什么", 1)
    add_para(
        doc,
        "Redis 在这个项目里主要有两个角色：一是缓存读多写少的数据，二是帮助处理幂等。缓存可以提升速度，但更新数据库后必须删除或刷新缓存，否则用户可能看到旧数据。",
    )
    add_table(
        doc,
        ["缓存键", "服务", "保存内容", "何时删除/更新"],
        [
            ("mall:auth:session:{sessionNo}", "auth-service", "登录会话快照", "登出、踢出会话、重置密码、角色状态变更"),
            ("mall:user:profile:{userId}", "user-service", "用户资料", "更新用户资料"),
            ("mall:user:addresses:{userId}", "user-service", "地址列表", "增删改地址、设置默认地址"),
            ("mall:product:*", "product-service", "分类、列表、详情、后台统计", "后台商品增删改后批量删除"),
            ("mall:cart:items:{userId}", "cart-service", "购物车列表", "加入、删除、下单清理"),
            ("mall:inventory:stock:{skuId}", "inventory-service", "库存快照", "预占、释放、扣减、补货、覆盖库存"),
            ("mall:order:submit:idempotency:{userId}:{key}", "order-service", "下单幂等锁或结果", "锁 10 分钟，成功结果缓存 1 天"),
            ("mall:order:detail:{orderNo}", "order-service", "订单详情", "取消、支付成功、发货、完成、删除"),
            ("mall:payment:detail:{paymentNo}", "payment-service", "支付详情", "支付成功、关闭支付单"),
        ],
    )

    add_heading(doc, "16. MySQL 表结构怎么理解", 1)
    add_para(doc, "每个服务都有自己的 schema.sql，启动时通过 MysqlSchemaInitializerSupport 自动尝试建库建表。核心表可以按业务域记忆：")
    add_table(
        doc,
        ["业务域", "数据库/表", "保存什么"],
        [
            ("认证", "mall_auth.auth_account", "登录名、密码哈希、状态、角色"),
            ("认证", "mall_auth.auth_token_session", "token 会话、jti、过期时间"),
            ("用户", "mall_user.ums_user", "用户昵称、手机、邮箱、角色"),
            ("用户", "mall_user.ums_user_address", "收货地址"),
            ("商品", "mall_product.pms_category / pms_spu / pms_sku", "分类、产品系列、具体 SKU"),
            ("购物车", "mall_cart.oms_cart_item", "用户购物车商品和数量"),
            ("库存", "mall_inventory.inventory_stock", "available_qty、locked_qty"),
            ("库存", "mall_inventory.inventory_reservation / item", "订单库存预占记录"),
            ("订单", "mall_order.oms_order / oms_order_item", "订单主表和商品明细"),
            ("订单", "mall_order.oms_order_submit_log", "下单幂等键和订单号对应关系"),
            ("支付", "mall_payment.pay_payment", "支付单、支付状态、第三方交易号"),
        ],
    )
    add_callout(
        doc,
        "为什么每个服务一个库",
        "微服务设计里推荐服务拥有自己的数据。订单服务不应该直接改库存表，库存服务也不应该直接改订单表。当前项目虽然都在本地 MySQL，但逻辑上已经按服务拆库。",
    )

    add_heading(doc, "17. 前端项目怎么工作", 1)
    add_para(
        doc,
        "mall-web 是 React + TypeScript + Vite 前端。main.tsx 挂载应用，App.tsx 配置页面路由，mallApi.ts 统一封装后端接口，AuthContext 统一保存登录状态，ToastContext 统一弹出提示。",
    )
    add_table(
        doc,
        ["前端文件", "作用"],
        [
            ("src/main.tsx", "React 入口，使用 BrowserRouter 包住 App"),
            ("src/App.tsx", "配置路由：首页、商品、购物车、结算、账号中心、后台"),
            ("src/api/client.ts", "封装 fetch、统一处理 CommonResponse、自动刷新 token"),
            ("src/api/mallApi.ts", "把后端接口包装成前端方法，如 submitOrder、createPayment"),
            ("src/context/AuthContext.tsx", "保存 session、profile，提供 login/register/logout"),
            ("src/context/ToastContext.tsx", "全局消息提示"),
            ("src/pages/HomePage.tsx", "首页商品推荐和加购物车"),
            ("src/pages/ProductPage.tsx", "商品详情、数量选择、加入购物车、去结算"),
            ("src/pages/CartPage.tsx", "购物车列表和删除商品"),
            ("src/pages/CheckoutPage.tsx", "从购物车生成订单，随后创建支付单"),
            ("src/pages/AdminPage.tsx", "后台管理商品、库存、订单、用户"),
        ],
    )
    add_heading(doc, "17.1 前端请求封装逻辑", 2)
    add_numbered(
        doc,
        [
            "request() 先从 localStorage 读取 accessToken。",
            "需要登录的接口如果没有 token，会抛出“请先登录”。",
            "fetch 返回后，parseEnvelope() 检查 HTTP 状态和后端 code。",
            "如果接口返回 401，并且本地有 refreshToken，会自动调用 /auth/token/refresh。",
            "刷新成功后，用新 accessToken 重试原请求。",
            "如果网络请求失败，部分商品接口会使用 fallback 数据，让前端仍能展示演示内容。",
        ],
    )

    add_heading(doc, "18. 配置与部署", 1)
    add_para(
        doc,
        "每个服务的 application.yml 都定义自己的端口、数据库、Redis 库编号、Nacos 开关和依赖服务 URL。默认 Nacos discovery/config 是 false，所以本地开发主要靠固定 localhost 地址。",
    )
    add_table(
        doc,
        ["配置项", "含义"],
        [
            ("MALL_VM_HOST", "中间件所在机器地址，默认 127.0.0.1"),
            ("MALL_MYSQL_PORT", "MySQL 端口，默认 3306"),
            ("MALL_MYSQL_USERNAME / PASSWORD", "MySQL 账号密码"),
            ("MALL_REDIS_PASSWORD", "Redis 密码"),
            ("MALL_AUTH_JWT_SECRET", "JWT 签名密钥，必须足够长"),
            ("MALL_AUTH_DEMO_PASSWORD", "演示账号密码"),
            ("MALL_NACOS_ENABLED", "是否启用 Nacos 服务发现，默认 false"),
        ],
    )
    add_para(doc, "deploy/docker-compose.yml 可以启动 MySQL、Redis、RocketMQ、Nacos 和 RocketMQ Dashboard。deploy/.env.example 是环境变量模板，不要把真实生产密码提交到 Git。")
    add_code(
        doc,
        r"""
# 启动全部本地后端服务
powershell -ExecutionPolicy Bypass -File .\deploy\start-local.ps1

# 同时启动后端和 Vite 前端开发服务
powershell -ExecutionPolicy Bypass -File .\deploy\start-dev.ps1 -SkipPackage

# 停止脚本启动的本地服务
powershell -ExecutionPolicy Bypass -File .\deploy\stop-local.ps1

# 前端单独开发
cd .\mall-web
npm run dev
""",
    )

    add_heading(doc, "19. 常见业务修改应该去哪改", 1)
    add_table(
        doc,
        ["需求", "主要修改位置", "提醒"],
        [
            ("新增商品字段", "product-service 的 pms_sku 表、DTO、ProductRepository、ProductService、前端类型和页面", "记得清理商品缓存"),
            ("修改订单状态", "OrderService 状态常量、状态判断、OrderRepository SQL、前端状态展示", "状态机不要随意跳转"),
            ("接入真实支付", "payment-service 新增渠道网关、回调验签、回调日志、对账任务", "不要让支付服务直接改订单表"),
            ("改登录策略", "AuthService、JwtTokenService、auth_account、auth_token_session", "注意旧 token 失效和 session 缓存"),
            ("增加库存流水", "inventory-service 新增 inventory_record 表和 Repository 写入", "预占、释放、扣减都要记流水"),
            ("接入 Nacos", "application.yml、服务启动配置、GatewayRoutesConfig 或服务发现路由", "默认当前代码用固定 URL"),
            ("用 RocketMQ 替代同步调用", "新增事件模型、生产者、消费者、消费幂等表", "先保证补偿和排障能力"),
        ],
    )

    add_heading(doc, "20. 代码阅读路线", 1)
    add_para(doc, "如果你是小白，不建议一上来全项目乱点。按下面路线读，会更顺。")
    add_numbered(
        doc,
        [
            "先读 README.md，知道怎么构建和启动。",
            "读根 pom.xml、mall-dependencies/pom.xml，理解模块和依赖版本。",
            "读 mall-common-core 的 CommonResponse、BusinessException，理解返回格式和错误处理。",
            "读 mall-gateway 的 GatewayRoutesConfig 和 GatewayAuthenticationFilter，理解请求怎么转发和鉴权。",
            "读 auth-service 的 AuthController、AuthService、JwtTokenService，理解登录和 token。",
            "读 product-service 的 ProductController、ProductService，理解商品如何展示。",
            "读 cart-service 的 CartService，理解加购物车。",
            "读 inventory-service 的 InventoryService，理解预占、释放、扣减。",
            "读 order-service 的 OrderService，重点看 submit、cancel、markPaid。",
            "读 payment-service 的 PaymentService，重点看 create、mockSuccess、closeByOrder。",
            "最后读 mall-web 的 api/client.ts、mallApi.ts、AuthContext.tsx，再看各个页面。",
        ],
    )

    add_heading(doc, "21. 新手调试建议", 1)
    add_bullets(
        doc,
        [
            "先确保 MySQL、Redis 已启动，且 .env 或系统环境变量里密码正确。",
            "优先通过网关访问接口，不要一开始就绕过网关，否则你可能拿不到 X-User-Id。",
            "遇到 401，先检查前端 localStorage 里有没有 accessToken，网关是否能访问 auth-service。",
            "遇到库存不足，查 inventory_stock 的 available_qty 和 locked_qty。",
            "遇到重复下单，查 Redis 幂等键和 oms_order_submit_log。",
            "遇到支付成功但订单没变，查 payment-service 调 order-service 是否报错，再查库存预占是否已释放。",
            "前端商品页能展示 fallback 数据，不代表后端一定启动了；要看网络请求是否真的成功。",
        ],
    )

    add_heading(doc, "22. 当前项目的边界和后续演进", 1)
    add_para(
        doc,
        "当前项目已经覆盖商城核心闭环，但仍是开发基线版本。它适合学习微服务拆分、登录鉴权、下单幂等、库存预占、支付回写这些关键概念。生产级系统还需要继续补强事件驱动、补偿任务、完整支付渠道、接口文档、自动化测试和监控告警。",
    )
    add_table(
        doc,
        ["方向", "当前情况", "建议演进"],
        [
            ("服务注册", "Nacos 配置已预留，默认关闭", "开启服务发现，网关路由从固定 URL 逐步变成服务名"),
            ("消息队列", "RocketMQ 已部署预留，主链路仍同步调用", "订单、支付、库存事件逐步改成 outbox + MQ"),
            ("支付", "mockSuccess 模拟支付成功", "接入真实支付宝/微信回调、验签、对账"),
            ("库存", "有预占/释放/扣减，但没有完整流水表", "增加库存流水和后台排障工具"),
            ("测试", "当前仓库未看到系统化测试", "补单元测试、集成测试、端到端测试"),
            ("API 文档", "docs 有 OpenAPI 清单", "生成可导入 Swagger/Knife4j 的 openapi.yaml"),
        ],
    )

    add_heading(doc, "附录 A：核心文件速查", 1)
    add_table(
        doc,
        ["文件", "看什么"],
        [
            (r"D:\ideaproject\mail\README.md", "项目说明、端口、启动命令"),
            (r"D:\ideaproject\mail\pom.xml", "根 Maven 模块"),
            (r"D:\ideaproject\mail\mall-dependencies\pom.xml", "依赖版本"),
            (r"D:\ideaproject\mail\mall-gateway\src\main\java\com\mall\gateway\filter\GatewayAuthenticationFilter.java", "网关鉴权"),
            (r"D:\ideaproject\mail\mall-services\auth-service\src\main\java\com\mall\auth\service\AuthService.java", "登录、注册、刷新、登出"),
            (r"D:\ideaproject\mail\mall-services\inventory-service\src\main\java\com\mall\inventory\service\InventoryService.java", "库存预占、释放、扣减"),
            (r"D:\ideaproject\mail\mall-services\order-service\src\main\java\com\mall\order\service\OrderService.java", "提交订单、取消、支付回写"),
            (r"D:\ideaproject\mail\mall-services\payment-service\src\main\java\com\mall\payment\service\PaymentService.java", "支付单创建和模拟支付成功"),
            (r"D:\ideaproject\mail\mall-web\src\api\client.ts", "前端请求封装和自动刷新 token"),
            (r"D:\ideaproject\mail\mall-web\src\api\mallApi.ts", "前端调用后端接口的统一入口"),
        ],
    )

    add_heading(doc, "附录 B：最小核心链路复习", 1)
    add_callout(
        doc,
        "从用户点击到数据库变化",
        "用户登录拿 token；网关验 token；商品服务展示 SKU；购物车服务保存临时商品；订单服务做幂等、查商品、锁库存、写订单；支付服务创建支付单；模拟支付成功后订单扣库存并变为待发货。这就是当前项目最重要的一条主线。",
    )

    return doc


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("商城微服务项目技术文档 - 小白友好版")
        run.font.size = Pt(8)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(MUTED)


def main() -> None:
    DESKTOP.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
